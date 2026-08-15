"""Integration tests for ExifTool commit pipeline (skipped if exiftool missing)."""

from __future__ import annotations

import json
import shutil
import subprocess
from pathlib import Path

import pytest

from engine import (
    _NO_WINDOW,
    ExifToolSession,
    _build_commit_args,
    _write_docx_metadata,
    _write_xlsx_metadata,
    _write_document_metadata,
    execute_commit,
    execute_commit_batch,
    log_commit_batch,
    list_undo_batches,
    rollback_last_batch,
    UNDO_LOG_FILE,
)

HAS_EXIFTOOL = shutil.which("exiftool") is not None
HAS_FFMPEG = shutil.which("ffmpeg") is not None


def _write_valid_media(path: Path, ext: str) -> None:
    """Create a real media file that ExifTool can process (dummy bytes can't)."""
    if not HAS_FFMPEG:
        path.write_bytes(b"\x00" * 1024)
        return
    cmd = ["ffmpeg", "-hide_banner", "-loglevel", "error", "-y"]
    if ext == ".mp4":
        cmd += ["-f", "lavfi", "-i", "color=red:size=64x64:duration=1",
                "-c:v", "libx264", "-pix_fmt", "yuv420p", str(path)]
    elif ext == ".jpg":
        cmd += ["-f", "lavfi", "-i", "color=blue:size=64x64", "-frames:v", "1",
                "-q:v", "2", str(path)]
    elif ext == ".png":
        cmd += ["-f", "lavfi", "-i", "color=blue:size=64x64", "-frames:v", "1", str(path)]
    else:
        path.write_bytes(b"\x00" * 1024)
        return
    subprocess.run(cmd, check=True, capture_output=True)


def _make_asset(tmp_path: Path, name: str = "test_video", ext: str = ".mp4",
                category: str = "test", valid_media: bool = False) -> dict:
    """Create a minimal asset dict for testing."""
    src = tmp_path / "src" / f"{name}{ext}"
    src.parent.mkdir(parents=True, exist_ok=True)
    if valid_media:
        _write_valid_media(src, ext)
    else:
        src.write_bytes(b"\x00" * 1024)
    return {
        "original_name": f"{name}{ext}",
        "original_path": src,
        "staged_name": f"renamed_{name}",
        "category": category,
        "tags": ["tag1", "tag2"],
        "summary": "Test summary for commit",
        "new_filename": f"renamed_{name}{ext}",
    }


# --- _build_commit_args tests (no ExifTool needed) ---

class TestBuildCommitArgs:
    def test_video_returns_args(self, tmp_path: Path):
        asset = _make_asset(tmp_path, ext=".mp4")
        target = tmp_path / "renamed_test_video.mp4"
        shutil.copy2(asset["original_path"], target)
        args = _build_commit_args(asset, target)
        assert len(args) > 0
        assert any("XMP-dc:Title" in a for a in args)
        assert any("QuickTime:Title" in a for a in args)
        assert args[-1] == str(target)

    def test_image_returns_args(self, tmp_path: Path):
        asset = _make_asset(tmp_path, ext=".jpg")
        target = tmp_path / "renamed_test_video.jpg"
        shutil.copy2(asset["original_path"], target)
        args = _build_commit_args(asset, target)
        assert len(args) > 0
        assert any("EXIF:XPTitle" in a for a in args)

    def test_docx_returns_empty(self, tmp_path: Path):
        asset = _make_asset(tmp_path, ext=".docx")
        target = tmp_path / "renamed_test_video.docx"
        shutil.copy2(asset["original_path"], target)
        args = _build_commit_args(asset, target)
        assert args == []

    def test_xlsx_returns_empty(self, tmp_path: Path):
        asset = _make_asset(tmp_path, ext=".xlsx")
        target = tmp_path / "renamed_test_video.xlsx"
        shutil.copy2(asset["original_path"], target)
        args = _build_commit_args(asset, target)
        assert args == []

    def test_txt_returns_empty(self, tmp_path: Path):
        asset = _make_asset(tmp_path, ext=".txt")
        target = tmp_path / "renamed_test_video.txt"
        shutil.copy2(asset["original_path"], target)
        args = _build_commit_args(asset, target)
        assert args == []

    def test_pdf_returns_args(self, tmp_path: Path):
        asset = _make_asset(tmp_path, ext=".pdf")
        target = tmp_path / "renamed_test_video.pdf"
        shutil.copy2(asset["original_path"], target)
        args = _build_commit_args(asset, target)
        assert len(args) > 0

    def test_mp3_returns_args(self, tmp_path: Path):
        asset = _make_asset(tmp_path, ext=".mp3")
        target = tmp_path / "renamed_test_video.mp3"
        shutil.copy2(asset["original_path"], target)
        args = _build_commit_args(asset, target)
        assert any("ID3:TIT2" in a for a in args)

    def test_wav_returns_args(self, tmp_path: Path):
        asset = _make_asset(tmp_path, ext=".wav")
        target = tmp_path / "renamed_test_video.wav"
        shutil.copy2(asset["original_path"], target)
        args = _build_commit_args(asset, target)
        assert any("XMP-dc:Title" in a for a in args)

    def test_m4a_returns_args(self, tmp_path: Path):
        asset = _make_asset(tmp_path, ext=".m4a")
        target = tmp_path / "renamed_test_video.m4a"
        shutil.copy2(asset["original_path"], target)
        args = _build_commit_args(asset, target)
        assert any("QuickTime:Title" in a for a in args)

    def test_flac_returns_args(self, tmp_path: Path):
        asset = _make_asset(tmp_path, ext=".flac")
        target = tmp_path / "renamed_test_video.flac"
        shutil.copy2(asset["original_path"], target)
        args = _build_commit_args(asset, target)
        assert any("XMP-dc:Title" in a for a in args)


# --- Audio hash and similarity tests ---

class TestNativeMetadataWriters:
    def test_write_docx_metadata(self, tmp_path: Path):
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")
        docx_file = tmp_path / "test.docx"
        doc = Document()
        doc.add_paragraph("hello")
        doc.save(str(docx_file))
        _write_docx_metadata(docx_file, "My Title", "My Summary", ["tag1", "tag2"])
        doc2 = Document(str(docx_file))
        assert doc2.core_properties.title == "My Title"
        assert doc2.core_properties.subject == "My Summary"
        assert "tag1" in doc2.core_properties.keywords

    def test_write_xlsx_metadata(self, tmp_path: Path):
        try:
            from openpyxl import Workbook
        except ImportError:
            pytest.skip("openpyxl not installed")
        xlsx_file = tmp_path / "test.xlsx"
        wb = Workbook()
        wb.save(str(xlsx_file))
        _write_xlsx_metadata(xlsx_file, "Sheet Title", "Sheet Summary", ["alpha", "beta"])
        wb2 = Workbook(str(xlsx_file)) if False else None
        from openpyxl import load_workbook
        wb2 = load_workbook(str(xlsx_file))
        assert wb2.properties.title == "Sheet Title"
        assert "alpha" in wb2.properties.keywords

    def test_write_document_metadata_txt_skips(self, tmp_path: Path):
        txt = tmp_path / "notes.txt"
        txt.write_text("content")
        asset = {"staged_name": "notes", "summary": "s", "tags": ["t"]}
        result = _write_document_metadata(txt, asset)
        assert result is False

    def test_write_document_metadata_docx_writes(self, tmp_path: Path):
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")
        docx_file = tmp_path / "report.docx"
        doc = Document()
        doc.add_paragraph("content")
        doc.save(str(docx_file))
        asset = {"staged_name": "report_final", "summary": "Report summary", "tags": ["report", "final"]}
        result = _write_document_metadata(docx_file, asset)
        assert result is True


# --- execute_commit tests (requires ExifTool for media, skips for docs) ---

@pytest.mark.skipif(not (HAS_EXIFTOOL and HAS_FFMPEG), reason="ExifTool or FFmpeg not installed")
class TestExecuteCommit:
    def test_commit_video_renames_and_writes_metadata(self, tmp_path: Path):
        asset = _make_asset(tmp_path, ext=".mp4", category="test_cat", valid_media=True)
        target_dir = tmp_path / "output"
        session = ExifToolSession()
        try:
            result = execute_commit(asset, target_dir, sort_into_folders=True, exiftool_session=session)
            assert isinstance(result, Path)
            assert "test_cat" in str(result)
            assert (target_dir / result).exists()
        finally:
            session.process.stdin.close()

    def test_commit_image_renames(self, tmp_path: Path):
        asset = _make_asset(tmp_path, ext=".jpg", category="photos", valid_media=True)
        target_dir = tmp_path / "output"
        session = ExifToolSession()
        try:
            result = execute_commit(asset, target_dir, sort_into_folders=False, exiftool_session=session)
            assert isinstance(result, Path)
            assert (target_dir / result).exists()
        finally:
            session.process.stdin.close()

    def test_commit_skip_metadata(self, tmp_path: Path):
        asset = _make_asset(tmp_path, ext=".mp4", category="test", valid_media=True)
        target_dir = tmp_path / "output"
        session = ExifToolSession()
        try:
            result = execute_commit(asset, target_dir, sort_into_folders=False, exiftool_session=session, skip_metadata=True)
            assert isinstance(result, Path)
            assert (target_dir / result).exists()
        finally:
            session.process.stdin.close()

    def test_commit_skip_rename_copies(self, tmp_path: Path):
        asset = _make_asset(tmp_path, ext=".mp4", category="test", valid_media=True)
        target_dir = tmp_path / "output"
        session = ExifToolSession()
        try:
            result = execute_commit(asset, target_dir, sort_into_folders=False, exiftool_session=session, skip_rename=True)
            assert isinstance(result, Path)
            assert result.exists()
            assert asset["original_path"].exists()
        finally:
            session.process.stdin.close()

    def test_commit_duplicate_name_appends_counter(self, tmp_path: Path):
        asset = _make_asset(tmp_path, ext=".mp4", category="test", valid_media=True)
        target_dir = tmp_path / "output" / "test"
        target_dir.mkdir(parents=True)
        (target_dir / "renamed_test_video.mp4").write_bytes(b"\x00" * 100)
        session = ExifToolSession()
        try:
            result = execute_commit(asset, target_dir.parent, sort_into_folders=True, exiftool_session=session)
            assert isinstance(result, Path)
            assert "_1" in str(result)
        finally:
            session.process.stdin.close()


# --- execute_commit_batch tests ---

@pytest.mark.skipif(not (HAS_EXIFTOOL and HAS_FFMPEG), reason="ExifTool or FFmpeg not installed")
class TestExecuteCommitBatch:
    def test_batch_multiple_assets(self, tmp_path: Path):
        assets = [
            _make_asset(tmp_path, name="vid1", ext=".mp4", category="cat_a", valid_media=True),
            _make_asset(tmp_path, name="vid2", ext=".mp4", category="cat_b", valid_media=True),
        ]
        target_dir = tmp_path / "output"
        session = ExifToolSession()
        try:
            results = execute_commit_batch(assets, target_dir, sort_into_folders=True, exiftool_session=session)
            assert len(results) == 2
            for r in results:
                assert isinstance(r, Path)
                assert (target_dir / r).exists()
        finally:
            session.process.stdin.close()

    def test_batch_mixed_formats(self, tmp_path: Path):
        assets = [
            _make_asset(tmp_path, name="video1", ext=".mp4", category="video", valid_media=True),
            _make_asset(tmp_path, name="image1", ext=".jpg", category="image", valid_media=True),
        ]
        target_dir = tmp_path / "output"
        session = ExifToolSession()
        try:
            results = execute_commit_batch(assets, target_dir, sort_into_folders=True, exiftool_session=session)
            assert len(results) == 2
            for r in results:
                assert isinstance(r, Path)
        finally:
            session.process.stdin.close()

    def test_batch_skip_metadata(self, tmp_path: Path):
        assets = [_make_asset(tmp_path, name="v1", ext=".mp4", category="test", valid_media=True)]
        target_dir = tmp_path / "output"
        session = ExifToolSession()
        try:
            results = execute_commit_batch(assets, target_dir, sort_into_folders=False, exiftool_session=session, skip_metadata=True)
            assert len(results) == 1
            assert isinstance(results[0], Path)
        finally:
            session.process.stdin.close()


# --- Undo / Rollback tests ---

class TestUndoRollback:
    def setup_method(self):
        if UNDO_LOG_FILE.exists():
            UNDO_LOG_FILE.unlink()

    def teardown_method(self):
        if UNDO_LOG_FILE.exists():
            UNDO_LOG_FILE.unlink()

    def test_log_commit_batch_writes_file(self, tmp_path: Path):
        records = [{"original_path": str(tmp_path / "a.mp4"), "new_path": str(tmp_path / "out" / "b.mp4"), "category": "test", "tags": ["t"]}]
        result = log_commit_batch("batch-001", str(tmp_path / "out"), records)
        assert result.exists()
        batches = list_undo_batches()
        assert len(batches) == 1
        assert batches[0]["batch_id"] == "batch-001"

    def test_list_undo_batches_returns_newest_first(self, tmp_path: Path):
        for i in range(3):
            log_commit_batch(f"batch-{i}", str(tmp_path / "out"), [])
        batches = list_undo_batches()
        assert len(batches) == 3
        assert batches[0]["batch_id"] == "batch-2"

    def test_rollback_last_batch_moves_files_back(self, tmp_path: Path):
        orig = tmp_path / "source" / "video.mp4"
        orig.parent.mkdir(parents=True)
        orig.write_bytes(b"\x00" * 512)
        dest = tmp_path / "dest" / "test" / "renamed_video.mp4"
        dest.parent.mkdir(parents=True)
        shutil.copy2(str(orig), str(dest))
        orig.unlink()
        records = [{"original_path": str(orig), "new_path": str(dest), "category": "test", "tags": ["t"]}]
        log_commit_batch("rollback-test", str(tmp_path / "dest"), records)
        result = rollback_last_batch()
        assert result["ok"] is True
        assert result["restored"] == 1
        assert orig.exists()
        assert not dest.exists()

    def test_rollback_empty_log_returns_error(self):
        result = rollback_last_batch()
        assert result["ok"] is False
        assert "No undo batches" in result["errors"][0]
